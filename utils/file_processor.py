"""
File processing utilities for Mindrian
Handles PDF, DOCX, text file extraction, and image detection
"""

import os
from typing import Optional, Tuple
from pathlib import Path

# === Image Support Constants ===
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic', '.heif'}
IMAGE_MIME_TYPES = {
    '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.png': 'image/png',
    '.gif': 'image/gif',
    '.webp': 'image/webp',
    '.heic': 'image/heic',
    '.heif': 'image/heif',
}


def is_image_file(file_name: str) -> bool:
    """Check if file is a supported image type."""
    return Path(file_name).suffix.lower() in IMAGE_EXTENSIONS


def get_image_mime_type(file_name: str) -> str:
    """Get MIME type for image file."""
    return IMAGE_MIME_TYPES.get(Path(file_name).suffix.lower(), 'image/jpeg')


def extract_text_from_pdf(file_path: str, max_pages: int = 50) -> Tuple[str, dict]:
    """
    Extract text content from a PDF file.

    Args:
        file_path: Path to the PDF file
        max_pages: Maximum number of pages to extract (default 50)

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        num_pages = len(reader.pages)

        text_parts = []
        pages_extracted = min(num_pages, max_pages)

        for i in range(pages_extracted):
            page = reader.pages[i]
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        # Truncate if too long (keep ~50k chars for context)
        if len(full_text) > 50000:
            full_text = full_text[:50000] + "\n\n[... content truncated ...]"

        metadata = {
            "type": "pdf",
            "total_pages": num_pages,
            "pages_extracted": pages_extracted,
            "char_count": len(full_text),
            "truncated": len(full_text) >= 50000
        }

        return full_text, metadata

    except Exception as e:
        return f"Error extracting PDF: {str(e)}", {"type": "pdf", "error": str(e)}


def extract_text_from_docx(file_path: str) -> Tuple[str, dict]:
    """
    Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    try:
        from docx import Document

        doc = Document(file_path)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        # Truncate if too long
        if len(full_text) > 50000:
            full_text = full_text[:50000] + "\n\n[... content truncated ...]"

        metadata = {
            "type": "docx",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "char_count": len(full_text),
            "truncated": len(full_text) >= 50000
        }

        return full_text, metadata

    except Exception as e:
        return f"Error extracting DOCX: {str(e)}", {"type": "docx", "error": str(e)}


def extract_text_from_txt(file_path: str) -> Tuple[str, dict]:
    """
    Read text content from a plain text file.

    Args:
        file_path: Path to the text file

    Returns:
        Tuple of (text_content, metadata_dict)
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Truncate if too long
        truncated = False
        if len(content) > 50000:
            content = content[:50000] + "\n\n[... content truncated ...]"
            truncated = True

        metadata = {
            "type": "text",
            "char_count": len(content),
            "line_count": content.count('\n') + 1,
            "truncated": truncated
        }

        return content, metadata

    except Exception as e:
        return f"Error reading text file: {str(e)}", {"type": "text", "error": str(e)}


def process_uploaded_file(file_path: str, file_name: str) -> Tuple[str, dict]:
    """
    Process an uploaded file and extract its text content.

    Args:
        file_path: Path to the uploaded file
        file_name: Original filename (used to determine type)

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    ext = Path(file_name).suffix.lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext in ['.txt', '.md', '.csv', '.json', '.py', '.js', '.html', '.css']:
        return extract_text_from_txt(file_path)
    else:
        return f"Unsupported file type: {ext}", {"type": "unsupported", "extension": ext}


def format_file_context(file_name: str, content: str, metadata: dict) -> str:
    """
    Format extracted file content for inclusion in conversation context.

    Args:
        file_name: Name of the file
        content: Extracted text content
        metadata: Metadata about the extraction

    Returns:
        Formatted context string
    """
    file_type = metadata.get("type", "unknown")

    header = f"\n\n---\n**UPLOADED FILE: {file_name}**\n"

    if file_type == "pdf":
        header += f"(PDF: {metadata.get('pages_extracted', '?')}/{metadata.get('total_pages', '?')} pages, {metadata.get('char_count', 0):,} characters)\n"
    elif file_type == "docx":
        header += f"(Word Document: {metadata.get('paragraphs', '?')} paragraphs, {metadata.get('char_count', 0):,} characters)\n"
    elif file_type == "text":
        header += f"(Text file: {metadata.get('line_count', '?')} lines, {metadata.get('char_count', 0):,} characters)\n"

    if metadata.get("truncated"):
        header += "**Note: Content was truncated due to length.**\n"

    if metadata.get("error"):
        return header + f"Error: {metadata['error']}\n---\n"

    return header + "---\n" + content + "\n---\n"
